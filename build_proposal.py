#!/usr/bin/env python3
"""Rewrite PhD-Proposal(final).docx with candidate-specific research content."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "PhD-Proposal-Sample-EN.docx"
OUTPUT = ROOT / "PhD-Proposal(final).docx"

# Paragraph index -> new text (None keeps paragraph unchanged)
CONTENT: dict[int, str | None] = {
    4: "Sajib Kumar Mitra",
    8: (
        "I would like to conduct research on deployable hybrid quantum–classical edge "
        "intelligence for sustainable IoT security and anomaly detection."
    ),
    9: (
        "Billions of Internet of Things (IoT) endpoints generate continuous telemetry that "
        "must be analysed under strict energy, latency, and trust constraints. Sustainable "
        "intelligent IoT systems are judged by network-level outcomes—energy efficiency, "
        "resilience, cross-tier interoperability, and long operational lifetimes. Classical "
        "edge artificial intelligence (AI) already enables near-sensor inference and "
        "reduces backhaul, yet high-dimensional streaming workloads, adversarial security "
        "analytics, and post-quantum transition pressures motivate selective use of quantum "
        "artificial intelligence (Quantum AI) in hybrid pipelines rather than cloud-only "
        "variational models.\n"
        "Quantum AI—the fusion of quantum computing with machine learning—is increasingly "
        "proposed for IoT security, kernel analytics, and combinatorial resource allocation. "
        "Contemporary QML toolchains such as PennyLane and Qiskit make hybrid "
        "quantum–classical learning pipelines practical for research, while edge gateways "
        "and FPGA SoCs create a realistic path toward compressed on-device inference. "
        "My ongoing systematic survey on Quantum AI-powered embedded devices for "
        "sustainable intelligent IoT systems (PRISMA-inspired corpus n=57) classifies IoT "
        "deployments into four paradigms—conventional, AI-based, quantum-enhanced, and "
        "quantum AI-based IoT—and shows that accuracy, circuit depth, algorithmic "
        "complexity, and deployable resource proxies are rarely reported jointly. This "
        "documented gap motivates doctoral research that co-designs hybrid models, "
        "benchmark protocols, and embedded prototypes under explicit power and latency "
        "budgets."
    ),
    11: "1-2. Research Objectives and Challenges",
    12: (
        "Hybrid quantum–classical edge intelligence for IoT anomaly detection requires "
        "co-design across machine learning models, quantum circuits, and resource-bounded "
        "embedded platforms. A deployable edge node or gateway must respect memory, "
        "latency, and energy constraints while sustaining detection quality under data "
        "sovereignty and trust policies. In the proposed research, the training phase uses "
        "simulators or cloud QPUs under qubit-depth constraints, while the edge phase "
        "exports compressed models to an FPGA bitstream or optimised TinyML runtime for "
        "local anomaly scoring on streaming features."
    ),
    13: (
        "Despite rapid progress in quantum machine learning (QML) for IoT, prior work "
        "remains largely simulator- or cloud-centric. For example, hybrid QML anomaly "
        "detection reports competitive accuracy but substantial simulator-era overhead; "
        "projected quantum kernels for IoT data analysis provide strong cloud-side "
        "analytics but limited edge deployability evidence; and recent embedded QML "
        "feasibility studies argue for hybrid architectures while offering sparse "
        "end-to-end deployment measurements. My four-paradigm survey further shows that "
        "smart-building IoT security gateways and post-discharge remote patient "
        "monitoring (RPM) in the Internet of Medical Things (IoMT) are representative "
        "use cases where hybrid edge analytics must satisfy zero-trust deployment "
        "constraints."
    ),
    14: (
        "Four barriers currently block practical edge adoption (Problem 1–4). First, "
        "cloud-centric training assumptions size variational quantum circuits (VQCs) for "
        "QPUs rather than microcontroller or FPGA budgets. Second, missing metric "
        "alignment means studies rarely report accuracy jointly with circuit depth, gate "
        "count, inference latency, and energy per decision on comparable IoT workloads. "
        "Third, the deployment path remains weak: few prototypes demonstrate "
        "train-on-cloud/simulator, infer-on-edge pipelines with open datasets and "
        "reproducible baselines. Fourth, over-broad quantum advantage claims seldom "
        "specify edge power caps, target false-alarm rates, or classical strong baselines."
    ),
    15: (
        "To address these challenges, the doctoral research aim is to design, implement, "
        "and empirically validate a deployable hybrid quantum–classical edge intelligence "
        "framework for IoT anomaly detection that meets predefined energy and latency "
        "budgets while matching or exceeding classical edge baselines within a stated "
        "accuracy tolerance. The work will extend the survey's PRISMA-style corpus and "
        "benchmark functional J(𝒲) into an open evaluation harness, develop a compressed "
        "hybrid model family, implement inference on one primary embedded platform "
        "(Xilinx Zynq/Artix-class FPGA or ARM IoT gateway with optional FPGA offload), "
        "and publish comparative evidence on when hybrid edge Quantum AI is justified "
        "versus classical edge AI alone."
    ),
    16: (
        "The research will investigate three questions. RQ1: Under a fixed edge power budget "
        "P and latency cap T, can compressed hybrid quantum–classical inference match "
        "classical edge accuracy within tolerance ε on standard IoT intrusion/anomaly "
        "datasets? RQ2: Which circuit compression strategies—depth reduction, angle "
        "encoding, and distillation to quantum-inspired classical models—minimise resource "
        "use without dominating accuracy loss? RQ3: Does FPGA or dedicated accelerator "
        "offload improve joules-per-inference versus pure software on ARM gateways for the "
        "proposed hybrid models?"
    ),
    19: (
        "My motivation for doctoral study stems from a desire to provide the world with "
        "highly secure and reliable information and communication systems, thereby "
        "contributing to a convenient and safe intelligent society. During my B.Sc. and M.S. "
        "studies in Computer Science and Engineering at the University of Dhaka, I became "
        "fascinated by the breadth and depth of security, quantum computing, and embedded "
        "systems. My M.S. thesis on efficient reversible logic and quantum circuit design, "
        "together with subsequent IEEE/ACM conference publications, established a research "
        "foundation in low-power quantum-classical co-design."
    ),
    20: (
        "As Joint Director (ICT) at Bangladesh Bank, I have gained extensive experience in "
        "secure ICT architecture, software engineering leadership, and national-scale "
        "systems relevant to information-system audit and trustworthy operations. The "
        "field of hybrid Quantum AI for IoT is a convergence of machine learning, quantum "
        "computing, cryptography, and embedded deployment; conceptual design, metric "
        "alignment, and reproducible evaluation are therefore crucial. I believe that my "
        "combination of academic research in reversible/quantum computing, industry-scale "
        "ICT practice, and co-investigator experience on a systematic IoT–Quantum AI "
        "survey will be an advantage in advancing this doctoral research. Through the "
        "program, I aim to cultivate research execution ability on hardware-aware QML "
        "pipelines and technical leadership that bridges academic research and "
        "technological development."
    ),
    24: (
        "This research adopts a hardware-aware two-phase pipeline rather than a full-stack "
        "“quantum IoT” platform. In the training phase, Qiskit and PennyLane on "
        "simulators or cloud QPUs support hyperparameter search under qubit-depth "
        "constraints, with optional post-quantum security context for threat models. In the "
        "edge phase, compressed models are exported to an FPGA bitstream or optimised "
        "C++/TinyML runtime on a Raspberry Pi-class or industrial gateway for local "
        "anomaly scoring. In the evaluation phase, unified metrics aligned with the "
        "survey benchmark table—algorithmic complexity class, wall-clock latency, energy "
        "proxy, and F1/AUC—are used to compare hybrid, quantum-inspired, and classical "
        "TinyML baselines on public IoT security datasets such as NSL-KDD, Edge-IIoTset, "
        "and Bot-IoT."
    ),
    25: (
        "The relationship between hybrid QML and trustworthy IoT deployment remains "
        "central even when quantum circuits are only a selective component of the analytics "
        "stack. Sustainable designs must satisfy deployability predicates on flash and RAM "
        "budgets, end-to-end latency service-level objectives, and policy constraints on "
        "data residency. Therefore, when implementing edge inference for IoT security "
        "analytics, it is crucial that model compression, benchmark reporting, and "
        "platform integration are implemented coherently rather than optimising accuracy "
        "alone. Reversible-logic optimisation may be explored optionally in Year 2 as a "
        "single arithmetic datapath block with switching-activity measurement, but not as "
        "a parallel thesis pillar."
    ),
    28: (
        "This research will proceed by repeating the following steps. Results will be "
        "presented at international conferences and journals—including IEEE Internet of "
        "Things Journal, Future Generation Computer Systems, Applied Soft Computing, IEEE "
        "Access, and selective venues such as IEEE IoT and QCE—and compiled into a doctoral "
        "dissertation with open-source artifacts."
    ),
    29: "<A. Benchmark Protocol and Technical Requirements>",
    30: (
        "First, define use cases and system models for hybrid edge IoT security analytics, "
        "including relevant stakeholders. Specifically, define smart-building IoT security "
        "gateway intrusion detection, industrial IoT anomaly detection, and post-discharge "
        "RPM IoMT scenarios from the perspective of edge platform operators, FPGA lab "
        "partners, IoT testbed operators, supervisors, and the open-source QML community. "
        "Clarify prerequisites and relationships among these actors, including dataset "
        "custodians, model trainers, and on-device inference services."
    ),
    31: (
        "Next, conduct a PRISMA-inspired literature update of IoT–Quantum AI research and "
        "formulate hypotheses for the deployability, latency, and security models required "
        "for the established system model. Analyze the technical requirements for benchmark "
        "protocols and hybrid model extensions by investigating prior QML for IoT anomaly "
        "detection, projected quantum kernels, embedded QML feasibility studies, and the "
        "candidate's four-paradigm survey corpus. Establish classical edge baselines (e.g., "
        "lightweight CNN or autoencoder) and implement a reproducible evaluation harness "
        "reporting accuracy, latency, algorithmic complexity, and energy proxy under fixed "
        "budgets. This step corresponds to Year 1 deliverables: benchmark specification, "
        "harness implementation, and initial classical results."
    ),
    32: "<B. Hybrid Model Development, Hardware Integration, and Verification>",
    33: (
        "Based on the aforementioned technical requirements, propose and verify compressed "
        "VQC, quantum-kernel, or quantum-inspired surrogate models with explicit qubit-depth "
        "and parameter budgets. Evaluate the validity of the deployability and security "
        "hypotheses through ablation studies on compression strategies (depth reduction, "
        "angle encoding, and distillation) and noise-aware training."
    ),
    34: (
        "Specifically, design the train–deploy split: simulator/cloud training, model "
        "compression and export, and edge inference on one primary platform (Zynq/Artix-class "
        "FPGA SoC or ARM gateway with optional FPGA offload). Integrate power measurement "
        "using a USB power meter or board PMIC, and compare head-to-head results versus "
        "classical TinyML and cloud-hybrid QML under fixed power and latency budgets to "
        "answer RQ1–RQ3. Peripheral integration includes public datasets, orchestration in "
        "Python/PyTorch, and embedded toolchains such as Vivado/Vitis HLS."
    ),
    35: (
        "The proposed methods will be evaluated for feasibility and scientific rigor. For "
        "feasibility, implement a prototype demonstrating train-on-simulator/cloud and "
        "infer-on-edge operation, with operational verification and performance "
        "measurement on the selected platform. For comparative evaluation, publish when "
        "hybrid edge Quantum AI is and is not justified under stated budgets. Expected "
        "contributions are: (i) a reference architecture for hybrid quantum–classical edge "
        "IoT anomaly detection with explicit deployability assumptions (TRL 2–4); (ii) an "
        "open benchmark protocol linking accuracy to circuit complexity and energy; and "
        "(iii) empirical decision guidance for practitioners. The 48-month work plan "
        "proceeds from literature update and baseline models (Year 1), hybrid VQC design "
        "and compression (Year 2), FPGA/embedded integration and large-scale evaluation "
        "(Year 3), to thesis writing, artifact release, and examination preparation (Year 4). "
        "Risks such as NISQ access/cost, scope creep, and negative results will be mitigated "
        "by prioritising simulators, constraining workload and platform scope, and treating "
        "null results as publishable evidence on when not to use hybrid QML."
    ),
    37: "4. Career paths after Doctoral",
    38: (
        "Technological innovation in IoT, edge AI, and quantum computing is accelerating, "
        "and in order to lead security and embedded-intelligence R&D in the future, I feel "
        "it is necessary to solidify my own area of expertise and to manage highly "
        "specialized research activities from a technical standpoint. After obtaining my "
        "degree, I aim to establish expertise as a researcher in deployable hybrid edge "
        "intelligence for IoT security."
    ),
    39: (
        "I also intend to leverage the research skills cultivated during doctoral study in "
        "corporate R&D management and technical guidance, becoming a person who can bridge "
        "the gap between academic research and technological development. Through these "
        "activities, I hope to collaborate with supervisors, FPGA and IoT testbed partners, "
        "and industry stakeholders to create synergies between foundational survey insights "
        "and reproducible prototypes, thereby increasing my contribution toward social "
        "implementation. Ultimately, I aspire to co-design trustworthy, energy-efficient "
        "IoT security analytics and to propose the architectural foundations needed for a "
        "convenient and secure intelligent society."
    ),
}


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text


def build() -> Path:
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)

    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))

    for idx, new_text in CONTENT.items():
        if new_text is None:
            continue
        if idx >= len(doc.paragraphs):
            raise IndexError(f"Paragraph index {idx} out of range ({len(doc.paragraphs)})")
        set_paragraph_text(doc.paragraphs[idx], new_text)

    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
